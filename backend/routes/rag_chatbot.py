from flask import Blueprint, request, jsonify, send_from_directory
from flask_jwt_extended import jwt_required, get_jwt_identity
from werkzeug.utils import secure_filename
import os
from datetime import datetime, timedelta
import csv
import re
from io import StringIO
from uuid import uuid4
from html import escape

# Ensure tracing is disabled
os.environ["LANGCHAIN_TRACING_V2"] = "false"

rag_bp = Blueprint('rag', __name__, url_prefix="/rag")
AI_PROMPT_KEY = "ai_assistant_system_prompt"
OPENAI_API_KEY_SETTING = "openai_api_key"
DEFAULT_AI_PROMPT = """You are GraceWise AI Coach: the voice of an experienced, calm homeschool parent and mentor—faith-centered, practical, and easy to talk to.

Who you are talking to: homeschooling parents (often tired, juggling multiple kids, and needing clarity fast). Speak with them, not at them.

How to sound natural:
- Use everyday language. Short sentences are fine. Avoid corporate or textbook tone.
- Match their energy: if they are stressed, be steady and kind; if they are chatty, you can be a little warmer and lighter.
- Do not start every reply with filler praise ("Great question!", "Absolutely!") unless it truly fits.
- Vary how you open and close. Do not end every message with the same catchphrase or Bible-adjacent sign-off.
- When Scripture or encouragement fits, keep it gentle and brief—not preachy.
- If they mix English with Roman Urdu or short Urdu phrases, understand intent and reply helpfully (you may stay in clear English unless they write fully in Urdu).

What you help with:
- Planning rhythms, subjects, pacing, and realistic schedules.
- Motivation, burnout, sibling days, and "what do I do tomorrow?"
- Gentle faith perspective when it helps—not forced into every answer.

Keep answers as long as they need to be: quick questions get concise answers; "help me think this through" can be longer with bullets or steps."""
COACH_RAG_SUPPLEMENT = """
When reference excerpts from the family's library are included in the user message:
- Use them when they clearly apply; blend facts into a natural reply without repeating "according to the context."
- If excerpts are off-topic or thin, say that lightly and still coach from experience—do not pretend the documents said something they did not."""
MANDATORY_CONTINUITY_RULES = """Conversation continuity rules:
- Always use prior turns from this same chat session as your primary conversation memory.
- Resolve follow-up references like "ye", "isko", "us table ko", "download link do" using the most recent relevant assistant output.
- If the user asks for a follow-up action (download/export/reformat/share), apply it to the last generated result unless they specify a different target.
- Ask a short clarification question only if more than one previous item could match the user's request.
- If your previous reply asked coaching/probing questions and the user responds, treat that as the answer to those exact questions first.
- Do not repeat the same options/questions after the user has answered; continue the same thread with the requested next step.
"""

# Configure upload settings
UPLOAD_FOLDER = 'documents'
EXPORT_FOLDER = os.path.join(UPLOAD_FOLDER, 'exports')
ALLOWED_EXTENSIONS = {'pdf', 'txt'}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
EXPORT_TTL_SECONDS = 3 * 60 * 60  # 3 hours

# Ensure documents folder exists
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
if not os.path.exists(EXPORT_FOLDER):
    os.makedirs(EXPORT_FOLDER)

# Global variables for lazy initialization
vector_data = None
llm_cache = None


def get_user_id():
    user_id = get_jwt_identity()
    if isinstance(user_id, str):
        user_id = int(user_id)
    return user_id


def is_admin_user():
    from models import User

    user = User.query.get(get_user_id())
    return bool(user and user.is_admin)


def get_ai_system_prompt():
    from models import AppSetting

    setting = AppSetting.query.filter_by(setting_key=AI_PROMPT_KEY).first()
    base_prompt = DEFAULT_AI_PROMPT
    if setting and setting.setting_value and setting.setting_value.strip():
        base_prompt = setting.setting_value

    merged = base_prompt.rstrip()
    if COACH_RAG_SUPPLEMENT.strip() not in merged:
        merged = f"{merged}\n{COACH_RAG_SUPPLEMENT}"

    if MANDATORY_CONTINUITY_RULES.strip() in merged:
        return merged

    return f"{merged}\n\n{MANDATORY_CONTINUITY_RULES}"


def _coach_llm_temperature():
    try:
        return float(os.environ.get("COACH_LLM_TEMPERATURE", "0.65"))
    except (TypeError, ValueError):
        return 0.65


def _build_retrieval_query_text(question, conversation_history, max_user_turns=3, max_chars=900):
    """
    Blend recent user turns with the latest message so embeddings catch follow-ups
    like "what about math?" or "same thing for my 7 year old".
    """
    user_chunks = []
    for msg in reversed(conversation_history or []):
        if msg.get("role") != "user":
            continue
        c = (msg.get("content") or "").strip()
        if c:
            user_chunks.append(c)
        if len(user_chunks) >= max_user_turns:
            break
    user_chunks.reverse()
    parts = user_chunks + [(question or "").strip()]
    combined = "\n".join(p for p in parts if p).strip()
    if len(combined) <= max_chars:
        return combined or (question or "").strip()
    return combined[-max_chars:]


def _parent_library_retrieve(question, conversation_history, chunks, chunk_embeddings, embed_model, top_k=6):
    """Semantic search over uploaded PDF chunks; returns (context_text, indices_used, weak_match)."""
    import numpy as np

    query_text = _build_retrieval_query_text(question, conversation_history)
    if not query_text.strip():
        return "", [], True

    query_emb = embed_model.encode([query_text])[0]
    qn = float(np.linalg.norm(query_emb))
    similarities = []
    for ce in chunk_embeddings:
        cn = float(np.linalg.norm(ce))
        if cn < 1e-10 or qn < 1e-10:
            similarities.append(0.0)
        else:
            similarities.append(float(np.dot(query_emb, ce) / (qn * cn)))

    ranked = sorted(range(len(similarities)), key=lambda i: similarities[i], reverse=True)
    top_idx = ranked[:top_k]
    if not top_idx:
        return "", [], True

    top_scores = [similarities[i] for i in top_idx[:3]]
    avg_top3 = sum(top_scores) / len(top_scores)
    weak_match = avg_top3 < 0.18

    context = "\n\n---\n\n".join(chunks[i] for i in top_idx)
    return context, top_idx, weak_match


def get_openai_api_key():
    from models import AppSetting

    setting = AppSetting.query.filter_by(setting_key=OPENAI_API_KEY_SETTING).first()
    if setting and setting.setting_value and setting.setting_value.strip():
        return setting.setting_value.strip()
    return os.environ.get("OPENAI_API_KEY")


def allowed_file(filename):
    """Check if file is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def documents_exist():
    for root, dirs, files in os.walk(UPLOAD_FOLDER):
        for file in files:
            if file.lower().endswith(".pdf"):
                return True
    return False


def _trim_title(text, max_len=60):
    text = (text or "").strip()
    if not text:
        return "New chat"
    return text[:max_len].strip() + ("..." if len(text) > max_len else "")


def _build_langchain_history(history):
    from langchain_core.messages import HumanMessage, AIMessage

    messages = []
    for msg in history:
        role = msg.get("role")
        content = (msg.get("content") or "").strip()
        if not content:
            continue
        if role == "user":
            messages.append(HumanMessage(content=content))
        elif role in ("assistant", "ai"):
            messages.append(AIMessage(content=content))
    return messages


def _has_download_intent(text):
    text = (text or "").strip().lower()
    if not text:
        return False
    action_keywords = [
        "download", "downloadable", "export", "file", "link",
        "create", "generate", "make", "save",
        "bna", "bana", "banado", "bna do", "bana do",
        "isko", "iska", "uska", "is table", "that table", "this table",
    ]
    format_pattern = re.compile(r"\b(csv|pdf|docx|txt)\b", re.IGNORECASE)
    return any(k in text for k in action_keywords) or bool(format_pattern.search(text))


def _split_markdown_row(row):
    row = row.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    return [cell.strip() for cell in row.split("|")]


def _is_separator_row(row):
    cells = _split_markdown_row(row)
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def _extract_latest_markdown_table(history):
    """
    Find the latest assistant markdown table from conversation history.
    Returns table lines or None.
    """
    for msg in reversed(history or []):
        if msg.get("role") not in ("assistant", "ai"):
            continue
        content = (msg.get("content") or "").strip()
        if not content:
            continue

        lines = content.splitlines()
        block = []
        latest_valid = None
        for line in lines:
            stripped = line.strip()
            if "|" in stripped and stripped.startswith("|"):
                block.append(stripped)
                continue
            if len(block) >= 2 and _is_separator_row(block[1]):
                latest_valid = block[:]
            block = []

        if len(block) >= 2 and _is_separator_row(block[1]):
            latest_valid = block[:]

        if latest_valid:
            return latest_valid
    return None


def _table_lines_to_csv(table_lines):
    if not table_lines or len(table_lines) < 2:
        return None

    rows = _normalize_markdown_table_rows(table_lines)
    if not rows:
        return None

    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(rows[0])
    writer.writerows(rows[1:])
    return output.getvalue()


def _normalize_markdown_table_rows(table_lines):
    if not table_lines or len(table_lines) < 2:
        return []

    header = _split_markdown_row(table_lines[0])
    if not header:
        return []

    col_count = len(header)
    rows = [header]

    for row_line in table_lines[2:]:
        cells = _split_markdown_row(row_line)
        if not any(cells):
            continue
        if len(cells) < col_count:
            cells.extend([""] * (col_count - len(cells)))
        elif len(cells) > col_count:
            cells = cells[:col_count]
        rows.append(cells)

    return rows


def _save_export_csv(csv_text):
    if not csv_text:
        return None

    filename = f"chat-table-{datetime.utcnow().strftime('%Y%m%d-%H%M%S')}-{uuid4().hex[:8]}.csv"
    file_path = os.path.join(EXPORT_FOLDER, filename)
    with open(file_path, "w", encoding="utf-8", newline="") as f:
        f.write(csv_text)
    return filename


def _save_export_text(text_content, extension="txt"):
    text_content = (text_content or "").strip()
    if not text_content:
        return None

    filename = f"chat-export-{datetime.utcnow().strftime('%Y%m%d-%H%M%S')}-{uuid4().hex[:8]}.{extension}"
    file_path = os.path.join(EXPORT_FOLDER, filename)
    with open(file_path, "w", encoding="utf-8", newline="\n") as f:
        f.write(text_content)
    return filename


def _export_expires_at_from_path(file_path):
    try:
        modified_at = datetime.utcfromtimestamp(os.path.getmtime(file_path))
    except OSError:
        return None
    return modified_at + timedelta(seconds=EXPORT_TTL_SECONDS)


def _is_export_expired(file_path):
    expires_at = _export_expires_at_from_path(file_path)
    if not expires_at:
        return False
    return datetime.utcnow() >= expires_at


def _cleanup_expired_exports():
    if not os.path.isdir(EXPORT_FOLDER):
        return 0

    removed = 0
    for name in os.listdir(EXPORT_FOLDER):
        path = os.path.join(EXPORT_FOLDER, name)
        if not os.path.isfile(path):
            continue
        if not _is_export_expired(path):
            continue
        try:
            os.remove(path)
            removed += 1
        except OSError:
            continue
    return removed


def _markdown_to_plain_text(markdown_text):
    blocks = _extract_markdown_blocks(markdown_text)
    rendered = []

    for block in blocks:
        if block["type"] == "text":
            text = _normalize_markdown_text_block(block["text"])
            if text:
                rendered.append(text)
        elif block["type"] == "table":
            ascii_table = _render_ascii_table(block["rows"])
            if ascii_table:
                rendered.append(ascii_table)

    return "\n\n".join(rendered).strip()


def _strip_markdown_inline(text):
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "").replace("__", "")
    return text


def _normalize_markdown_text_block(text):
    cleaned_lines = []
    for raw_line in (text or "").splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            cleaned_lines.append("")
            continue

        line = re.sub(r"^#{1,6}\s*", "", line)
        line = re.sub(r"^>\s*", "", line)
        line = _strip_markdown_inline(line)
        cleaned_lines.append(line.strip())

    normalized = "\n".join(cleaned_lines).strip()
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized


def _extract_markdown_blocks(markdown_text):
    lines = (markdown_text or "").splitlines()
    if not lines:
        return []

    blocks = []
    text_buffer = []
    idx = 0

    def flush_text():
        if not text_buffer:
            return
        combined = "\n".join(text_buffer).strip()
        if combined:
            blocks.append({"type": "text", "text": combined})
        text_buffer.clear()

    while idx < len(lines):
        line = lines[idx].strip()
        if (
            line.startswith("|")
            and idx + 1 < len(lines)
            and _is_separator_row(lines[idx + 1].strip())
        ):
            flush_text()
            table_lines = [lines[idx].strip(), lines[idx + 1].strip()]
            idx += 2

            while idx < len(lines):
                next_line = lines[idx].strip()
                if next_line.startswith("|"):
                    table_lines.append(next_line)
                    idx += 1
                else:
                    break

            rows = _normalize_markdown_table_rows(table_lines)
            if rows:
                blocks.append({"type": "table", "rows": rows})
            continue

        text_buffer.append(lines[idx])
        idx += 1

    flush_text()
    return blocks


def _render_ascii_table(rows):
    if not rows:
        return ""

    normalized_rows = []
    col_count = max(len(row) for row in rows)
    for row in rows:
        normalized = [(_strip_markdown_inline(str(cell or ""))).replace("\n", " ").strip() for cell in row]
        if len(normalized) < col_count:
            normalized.extend([""] * (col_count - len(normalized)))
        normalized_rows.append(normalized)

    col_widths = [0] * col_count
    for row in normalized_rows:
        for idx, cell in enumerate(row):
            col_widths[idx] = max(col_widths[idx], len(cell))

    def format_row(row):
        return "| " + " | ".join(cell.ljust(col_widths[idx]) for idx, cell in enumerate(row)) + " |"

    separator = "+-" + "-+-".join("-" * width for width in col_widths) + "-+"
    output_lines = [separator, format_row(normalized_rows[0]), separator]
    for row in normalized_rows[1:]:
        output_lines.append(format_row(row))
    output_lines.append(separator)
    return "\n".join(output_lines)


def _reportlab_inline_markup(text):
    """Limited markdown inline → ReportLab Paragraph XML (escape first)."""
    t = escape(text or "")
    t = re.sub(r"`([^`]+)`", r'<font face="Courier" size="9">\1</font>', t)
    t = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", t)
    t = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<i>\1</i>", t)
    return t.replace("\n", "<br/>")


def _build_pdf_styles():
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT

    base = getSampleStyleSheet()
    body = ParagraphStyle(
        "GWBody",
        parent=base["BodyText"],
        fontName="Helvetica",
        fontSize=11,
        leading=15,
        spaceAfter=8,
        alignment=TA_LEFT,
    )
    label_you = ParagraphStyle(
        "GWYou",
        parent=base["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        textColor=base["BodyText"].textColor,
        spaceBefore=14,
        spaceAfter=6,
        leading=14,
    )
    label_coach = ParagraphStyle(
        "GWCoach",
        parent=base["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        textColor=base["BodyText"].textColor,
        spaceBefore=14,
        spaceAfter=6,
        leading=14,
    )
    h3 = ParagraphStyle(
        "GWH3",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        spaceBefore=8,
        spaceAfter=4,
    )
    table_header = ParagraphStyle(
        "GWTableHead",
        parent=body,
        fontName="Helvetica-Bold",
        leading=13,
        fontSize=10,
    )
    return {"body": body, "label_you": label_you, "label_coach": label_coach, "h3": h3, "table_header": table_header}


def _pdf_append_rl_table(story, rows, doc_width, body_style, table_header_style):
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.units import inch

    if not rows:
        return
    col_count = len(rows[0]) if rows[0] else 1
    table_data = []
    for row_index, row in enumerate(rows):
        pstyle = table_header_style if row_index == 0 else body_style
        table_data.append([
            Paragraph(escape(str(cell or "")).replace("\n", "<br/>"), pstyle)
            for cell in row
        ])
    col_w = doc_width / max(col_count, 1)
    tbl = Table(table_data, repeatRows=1, colWidths=[col_w] * col_count)
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f2ecdf")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#3d3220")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d3c4a3")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fbf8f2")]),
            ]
        )
    )
    story.append(tbl)
    story.append(Spacer(1, 0.14 * inch))


def _pdf_append_formatted_body(story, text_block, doc_width, styles):
    """Format a body: ### headings, bullets, numbered lists, paragraphs; nested tables via blocks."""
    from reportlab.platypus import Paragraph, Spacer, ListFlowable, ListItem
    from reportlab.lib.units import inch

    if not (text_block or "").strip():
        return

    for block in _extract_markdown_blocks(text_block.strip()):
        if block["type"] == "table":
            _pdf_append_rl_table(
                story,
                block["rows"],
                doc_width,
                styles["body"],
                styles["table_header"],
            )
            continue

        raw = block["text"]
        lines = raw.splitlines()
        i = 0
        while i < len(lines):
            stripped = lines[i].strip()
            if not stripped:
                story.append(Spacer(1, 0.05 * inch))
                i += 1
                continue
            if stripped.startswith("### "):
                story.append(Paragraph(_reportlab_inline_markup(stripped[4:]), styles["h3"]))
                i += 1
                continue
            if re.match(r"^[-*]\s+", stripped):
                items = []
                while i < len(lines):
                    s = lines[i].strip()
                    if re.match(r"^[-*]\s+", s):
                        items.append(re.sub(r"^[-*]\s+", "", s))
                        i += 1
                    else:
                        break
                list_items = [
                    ListItem(
                        Paragraph(_reportlab_inline_markup(t), styles["body"]),
                        leftIndent=12,
                        bulletDedent="auto",
                    )
                    for t in items
                    if t.strip()
                ]
                if list_items:
                    story.append(
                        ListFlowable(
                            list_items,
                            bulletType="bullet",
                            bulletFontSize=9,
                            leftPadding=18,
                        )
                    )
                    story.append(Spacer(1, 0.08 * inch))
                continue
            if re.match(r"^\d+\.\s+", stripped):
                items = []
                while i < len(lines):
                    s = lines[i].strip()
                    if re.match(r"^\d+\.\s+", s):
                        items.append(re.sub(r"^\d+\.\s+", "", s))
                        i += 1
                    else:
                        break
                list_items = [
                    ListItem(
                        Paragraph(_reportlab_inline_markup(t), styles["body"]),
                        leftIndent=14,
                        bulletDedent="auto",
                    )
                    for t in items
                    if t.strip()
                ]
                if list_items:
                    story.append(
                        ListFlowable(
                            list_items,
                            bulletType="1",
                            leftPadding=22,
                        )
                    )
                    story.append(Spacer(1, 0.08 * inch))
                continue
            para_lines = []
            while i < len(lines):
                s = lines[i].strip()
                if not s:
                    break
                if s.startswith("### ") or re.match(r"^[-*]\s+", s) or re.match(r"^\d+\.\s+", s):
                    break
                para_lines.append(s)
                i += 1
            if para_lines:
                merged = " ".join(para_lines)
                story.append(Paragraph(_reportlab_inline_markup(merged), styles["body"]))
                story.append(Spacer(1, 0.06 * inch))


def _pdf_story_from_markdown(markdown_text, doc_width):
    """Build flowables: chat transcript (## You / ## AI Coach) or plain markdown document."""
    from reportlab.platypus import Paragraph, Spacer
    from reportlab.lib.units import inch

    story = []
    styles = _build_pdf_styles()
    md = (markdown_text or "").strip()
    if not md:
        return story

    if re.search(r"(?m)^## (You|AI Coach)\s*\n", md):
        chunks = re.split(r"(?m)^## ([^\n]+)\n", md)
        preamble = (chunks[0] or "").strip()
        if preamble:
            _pdf_append_formatted_body(story, preamble, doc_width, styles)
        it = iter(chunks[1:])
        for title, body in zip(it, it):
            t = (title or "").strip()
            b = (body or "").strip()
            tl = t.lower()
            if tl == "you":
                story.append(Paragraph("You", styles["label_you"]))
            elif tl == "ai coach":
                story.append(Paragraph("AI Coach", styles["label_coach"]))
            else:
                story.append(Paragraph(_reportlab_inline_markup(t), styles["h3"]))
            story.append(Spacer(1, 0.04 * inch))
            _pdf_append_formatted_body(story, b, doc_width, styles)
            story.append(Spacer(1, 0.1 * inch))
        return story

    for block in _extract_markdown_blocks(md):
        if block["type"] == "table":
            _pdf_append_rl_table(
                story,
                block["rows"],
                doc_width,
                styles["body"],
                styles["table_header"],
            )
        else:
            _pdf_append_formatted_body(story, block["text"], doc_width, styles)
    return story


def _save_export_pdf(markdown_text):
    md = (markdown_text or "").strip()
    if not md:
        return None, "No content available to export."

    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate
    except Exception:
        return None, "PDF export dependency is missing. Please install reportlab."

    filename = f"chat-export-{datetime.utcnow().strftime('%Y%m%d-%H%M%S')}-{uuid4().hex[:8]}.pdf"
    file_path = os.path.join(EXPORT_FOLDER, filename)

    try:
        doc = SimpleDocTemplate(
            file_path,
            pagesize=LETTER,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )
        story = _pdf_story_from_markdown(md, doc.width)
        if not story:
            return None, "No content available to export."

        doc.build(story)
        return filename, None
    except Exception as exc:
        return None, str(exc)


def _docx_flush_buffer(doc, buffer_lines):
    if not buffer_lines:
        return
    doc.add_paragraph(" ".join(buffer_lines))


def _docx_append_text_body(doc, text_block):
    """Headings, bullets, numbered lists, plain paragraphs (Word)."""
    lines = (text_block or "").splitlines()
    buf = []
    i = 0
    while i < len(lines):
        raw = lines[i]
        s = raw.strip()
        if not s:
            _docx_flush_buffer(doc, buf)
            buf = []
            i += 1
            continue
        if s.startswith("### "):
            _docx_flush_buffer(doc, buf)
            buf = []
            doc.add_heading(s[4:].strip(), level=3)
            i += 1
            continue
        if re.match(r"^[-*]\s+", s):
            _docx_flush_buffer(doc, buf)
            buf = []
            while i < len(lines):
                ls = lines[i].strip()
                if re.match(r"^[-*]\s+", ls):
                    doc.add_paragraph(re.sub(r"^[-*]\s+", "", ls), style="List Bullet")
                    i += 1
                else:
                    break
            continue
        if re.match(r"^\d+\.\s+", s):
            _docx_flush_buffer(doc, buf)
            buf = []
            while i < len(lines):
                ls = lines[i].strip()
                if re.match(r"^\d+\.\s+", ls):
                    doc.add_paragraph(re.sub(r"^\d+\.\s+", "", ls), style="List Number")
                    i += 1
                else:
                    break
            continue
        buf.append(s)
        i += 1
    _docx_flush_buffer(doc, buf)


def _docx_append_markdown_chunk(doc, md_text):
    for block in _extract_markdown_blocks((md_text or "").strip()):
        if block["type"] == "table":
            rows = block["rows"]
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for row_index, row in enumerate(rows):
                for col_index, cell in enumerate(row):
                    table_cell = table.rows[row_index].cells[col_index]
                    table_cell.text = str(cell or "")
                    if row_index == 0 and table_cell.paragraphs and table_cell.paragraphs[0].runs:
                        table_cell.paragraphs[0].runs[0].bold = True
            doc.add_paragraph("")
            continue
        _docx_append_text_body(doc, block["text"])


def _save_export_docx(markdown_text):
    md = (markdown_text or "").strip()
    if not md:
        return None, "No content available to export."

    try:
        from docx import Document
    except Exception:
        return None, "Word export dependency is missing. Please install python-docx."

    filename = f"chat-export-{datetime.utcnow().strftime('%Y%m%d-%H%M%S')}-{uuid4().hex[:8]}.docx"
    file_path = os.path.join(EXPORT_FOLDER, filename)

    try:
        doc = Document()

        if re.search(r"(?m)^## (You|AI Coach)\s*\n", md):
            chunks = re.split(r"(?m)^## ([^\n]+)\n", md)
            preamble = (chunks[0] or "").strip()
            if preamble:
                _docx_append_markdown_chunk(doc, preamble)
            it = iter(chunks[1:])
            for title, body in zip(it, it):
                doc.add_heading((title or "").strip(), level=2)
                _docx_append_markdown_chunk(doc, (body or "").strip())
        else:
            _docx_append_markdown_chunk(doc, md)

        doc.save(file_path)
        return filename, None
    except Exception as exc:
        return None, str(exc)


def _detect_requested_export_format(text):
    text = (text or "").lower()
    if not text:
        return None

    format_keywords = {
        "docx": ["docx", "word document", "word file", "ms word", "word", ".docx", ".doc"],
        "pdf": ["pdf", ".pdf"],
        "csv": ["csv", "excel", ".csv"],
        "txt": ["txt", "text file", ".txt"],
    }

    matches = []
    for export_format, keywords in format_keywords.items():
        indices = [text.find(k) for k in keywords if text.find(k) != -1]
        if indices:
            matches.append((min(indices), export_format))

    if not matches:
        return None
    matches.sort(key=lambda item: item[0])
    return matches[0][1]


def _extract_latest_assistant_content(history):
    for msg in reversed(history or []):
        if msg.get("role") not in ("assistant", "ai"):
            continue
        content = (msg.get("content") or "").strip()
        if content:
            return content
    return None


def _is_assistant_export_delivery_message(content):
    """True for auto-generated 'your file is ready' assistant bubbles (not real chat)."""
    if not content:
        return False
    low = content.lower()
    if "your file is ready" in low:
        return True
    if "download-expiry" in low or "data-expires-at" in low:
        return True
    if "/rag/exports/" in content or "/api/rag/exports/" in content:
        if "download" in low or "backup" in low:
            return True
    return False


def _conversation_transcript_markdown(history, max_messages=100):
    """
    Full chat as markdown: ## You / ## AI Coach sections (chronological).
    Omits export-stub assistant messages so PDF matches the real conversation.
    """
    if not history:
        return ""
    slice_hist = history[-max_messages:]
    parts = []
    for msg in slice_hist:
        role = (msg.get("role") or "").lower()
        content = (msg.get("content") or "").strip()
        if not content:
            continue
        if role in ("assistant", "ai"):
            if _is_assistant_export_delivery_message(content):
                continue
            parts.append(f"## AI Coach\n\n{content}")
        elif role == "user":
            parts.append(f"## You\n\n{content}")
    return "\n\n".join(parts).strip()


def _prose_markdown_for_document_export(history):
    """Prefer full transcript; fall back to latest assistant reply."""
    transcript = _conversation_transcript_markdown(history)
    if len(transcript.strip()) >= 40:
        return transcript
    latest = _extract_latest_exportable_assistant_content(history)
    if latest and len(latest.strip()) >= 10:
        return latest
    raw = _extract_latest_assistant_content(history)
    if raw and not _is_assistant_export_delivery_message(raw):
        return raw.strip()
    return transcript


def _is_clarifying_prompt_line(line):
    text = (line or "").strip().lower()
    if not text:
        return False

    normalized = text.rstrip(":")
    if normalized in ("coaching questions", "probing questions", "clarifying questions"):
        return True

    starters = (
        "thanks for clarifying",
        "before i put this together",
        "before i create this",
        "before i create that",
        "to tailor this",
        "to tailor that",
        "for example",
    )
    return any(text.startswith(prefix) for prefix in starters)


def _is_question_like_line(line):
    text = (line or "").strip()
    if not text:
        return False
    text = re.sub(r"^[-*+]\s*", "", text)
    text = re.sub(r"^\d+[\).\:-]\s*", "", text)
    return text.endswith("?")


def _sanitize_export_markdown(markdown_text):
    lines = (markdown_text or "").splitlines()
    if not lines:
        return ""

    output = []
    idx = 0

    while idx < len(lines):
        current = lines[idx]
        stripped = current.strip()

        if _is_clarifying_prompt_line(current):
            idx += 1
            while idx < len(lines):
                probe = lines[idx]
                probe_stripped = probe.strip()
                if not probe_stripped:
                    idx += 1
                    continue
                if _is_clarifying_prompt_line(probe) or _is_question_like_line(probe):
                    idx += 1
                    continue
                # Keep skipping bullet/numbered items that are written as questions.
                if re.match(r"^[-*+]\s+", probe_stripped) and "?" in probe_stripped:
                    idx += 1
                    continue
                if re.match(r"^\d+[\).\:-]\s+", probe_stripped) and "?" in probe_stripped:
                    idx += 1
                    continue
                break

            while output and not output[-1].strip():
                output.pop()
            continue

        output.append(current)
        idx += 1

    cleaned = "\n".join(output).strip()
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned


def _extract_latest_exportable_assistant_content(history):
    for msg in reversed(history or []):
        if msg.get("role") not in ("assistant", "ai"):
            continue
        content = (msg.get("content") or "").strip()
        if not content:
            continue
        sanitized = _sanitize_export_markdown(content)
        if sanitized:
            return sanitized
    return None


def _extract_questions_from_text(text, max_questions=3):
    text = (text or "").strip()
    if not text:
        return []

    questions = []
    seen = set()
    lines = text.splitlines()

    for line in lines:
        cleaned = re.sub(r"^\s*[-*+]\s*", "", line.strip())
        cleaned = re.sub(r"^\s*\d+[\).\:-]\s*", "", cleaned)
        if "?" not in cleaned:
            continue

        # Normalize likely encoding artifacts like "family?s" where "?" appears inside a word.
        normalized = re.sub(r"(?<=\w)\?(?=\w)", "'", cleaned)

        if normalized.count("?") <= 1:
            parts = [normalized]
        else:
            parts = [f"{part.strip()}?" for part in normalized.split("?") if part.strip()]

        for part in parts:
            question = " ".join(part.split())
            if len(question) < 4 or question in seen:
                continue
            seen.add(question)
            questions.append(question)
            if len(questions) >= max_questions:
                return questions

    if questions:
        return questions

    prose_parts = re.findall(r"[^?]{3,}\?", text)
    for part in prose_parts:
        question = " ".join(part.split())
        if len(question) < 4 or question in seen:
            continue
        seen.add(question)
        questions.append(question)
        if len(questions) >= max_questions:
            break
    return questions


def _extract_recent_assistant_questions(history, max_questions=3, assistant_turns=3):
    questions = []
    remaining_turns = assistant_turns

    for msg in reversed(history or []):
        if remaining_turns <= 0 or len(questions) >= max_questions:
            break
        if msg.get("role") not in ("assistant", "ai"):
            continue

        remaining_turns -= 1
        content = (msg.get("content") or "").strip()
        if not content:
            continue

        fresh = _extract_questions_from_text(content, max_questions=max_questions)
        for q in fresh:
            if q not in questions:
                questions.append(q)
            if len(questions) >= max_questions:
                break

    return questions


def _is_followup_answer_style(question):
    text = (question or "").strip().lower()
    if not text:
        return False

    markers = [
        "yes",
        "no",
        "both",
        "either",
        "first",
        "second",
        "weekly rhythm",
        "reading",
        "spelling",
        "support",
        "doable",
        "let's",
        "i would like",
        "i'd like",
        "start with",
        "focus on",
        "for now",
    ]
    if any(marker in text for marker in markers):
        return True

    # Most follow-up answers are short statements without a new question.
    if "?" not in text and len(text) <= 500:
        return True

    return False


def _build_followup_anchor_message(history, question):
    questions = _extract_recent_assistant_questions(history, max_questions=3, assistant_turns=1)
    if not questions:
        return None
    if not _is_followup_answer_style(question):
        return None

    rendered_questions = "\n".join(f"{idx + 1}. {q}" for idx, q in enumerate(questions))
    return (
        "Continuity (natural follow-up):\n"
        "Their latest message is probably answering something you asked earlier in this chat.\n"
        f"You had asked:\n{rendered_questions}\n\n"
        "Respond like a real coach: acknowledge what they picked or said, then move the conversation forward with clear next steps. "
        "Do not re-list the same choices unless they seem lost."
    )


def _download_link_label(filename):
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
    if ext == "csv":
        return "Download the CSV"
    if ext == "pdf":
        return "Download the PDF"
    if ext in ("doc", "docx"):
        return "Download the Word Document"
    return "Download the Document"


def _export_ready_chat_message(export_note, export_format, api_link, filename):
    """Single friendly download (no backup link), with a small card for the button + timer."""
    blurbs = {
        "pdf": "I saved our conversation as a **PDF** — easy to read later or print.",
        "docx": "I put everything into a **Word document** so you can edit it if you like.",
        "csv": "Here is your **spreadsheet (CSV)** — handy for sorting or sharing that table.",
        "txt": "Here is a **plain text** copy of the chat if you want something simple.",
    }
    blurb = blurbs.get(export_format, "Here is the file you asked for.")
    label = _download_link_label(filename)
    expires_at = (datetime.utcnow() + timedelta(seconds=EXPORT_TTL_SECONDS)).replace(microsecond=0).isoformat() + "Z"
    href = escape(api_link, quote=True)
    safe_label = escape(label)
    countdown = (
        f'<div class="download-expiry" data-expires-at="{escape(expires_at)}">'
        "Time remaining: calculating…"
        "</div>"
    )
    card = (
        '<div class="chat-download-card">'
        f'<p class="chat-download-hint">Tap when you are ready — the link stays good for a few hours, then you can ask me to export again.</p>'
        f"{countdown}"
        f'<a href="{href}" class="download-link chat-download-button" target="_blank" rel="noopener noreferrer">{safe_label}</a>'
        "</div>"
    )
    prefix = (export_note or "").strip()
    if prefix:
        prefix = prefix + "\n\n"
    return prefix + blurb + "\n\n" + card


def _build_export_links(filename):
    base = (request.headers.get("Origin") or request.url_root or "").rstrip("/")
    if not base:
        return {
            "api_link": f"/api/rag/exports/{filename}",
            "direct_link": f"/rag/exports/{filename}",
        }
    return {
        "api_link": f"{base}/api/rag/exports/{filename}",
        "direct_link": f"{base}/rag/exports/{filename}",
    }


def _try_auto_export(question, history):
    """
    If user asks for download/export, generate a file.
    - CSV: latest assistant markdown table (if any).
    - PDF / Word / TXT: full chat transcript when possible (You + AI Coach), else latest assistant reply.
    """
    _cleanup_expired_exports()

    if not _has_download_intent(question):
        return None

    latest_raw = _extract_latest_assistant_content(history)
    latest_content = _extract_latest_exportable_assistant_content(history)
    prose_md = _prose_markdown_for_document_export(history)

    requested_format = _detect_requested_export_format(question)
    table_lines = _extract_latest_markdown_table([{"role": "assistant", "content": latest_raw or ""}])
    export_format = requested_format or ("csv" if table_lines else "txt")

    filename = None
    export_note = ""
    export_error = None

    if export_format == "csv":
        if not table_lines:
            export_format = "txt"
            export_note = "I could not find a table, so I exported the conversation as a text document.\n\n"
        else:
            csv_text = _table_lines_to_csv(table_lines)
            filename = _save_export_csv(csv_text)
    elif export_format == "pdf":
        if not (prose_md or "").strip():
            return {
                "answer": "I could not find conversation content to export yet. Chat a bit first, then ask for the PDF."
            }
        filename, export_error = _save_export_pdf(prose_md)
    elif export_format == "docx":
        if not (prose_md or "").strip():
            return {
                "answer": "I could not find conversation content to export yet. Chat a bit first, then ask for the Word file."
            }
        filename, export_error = _save_export_docx(prose_md)
    else:
        if not (prose_md or "").strip():
            return {
                "answer": "I could not find conversation content to export yet. Chat a bit first, then ask for the file."
            }
        filename = _save_export_text(_markdown_to_plain_text(prose_md), extension="txt")
        export_format = "txt"

    if not filename:
        if export_error:
            return {"answer": f"I could not generate that file yet. {export_error}"}
        return None

    links = _build_export_links(filename)
    expires_at = (datetime.utcnow() + timedelta(seconds=EXPORT_TTL_SECONDS)).replace(microsecond=0).isoformat() + "Z"
    answer = _export_ready_chat_message(export_note, export_format, links["api_link"], filename)

    return {
        "answer": answer,
        "export_filename": filename,
        "export_format": export_format,
        "download_url": links["api_link"],
        "expires_at": expires_at,
        "ttl_seconds": EXPORT_TTL_SECONDS,
    }


def _prepare_chat_state(question, data, expected_chat_type=None, expected_lesson_id=None):
    """
    Prepare session/history for normal send and edited-message regeneration.
    Returns: (session, model_history, persist_state, error_dict_or_none)
    """
    from models import db, AiSession, AiChatMessage

    user_id = get_user_id()
    fallback_history = data.get("history", [])
    session_id = data.get("session_id")
    edit_message_id = data.get("edit_message_id")

    if not user_id or not session_id:
        return None, fallback_history, None, None

    try:
        session_id = int(session_id)
    except (TypeError, ValueError):
        return None, fallback_history, None, {"answer": "Invalid session_id."}

    session = AiSession.query.filter_by(id=session_id, user_id=user_id).first()
    if not session:
        return None, fallback_history, None, {"answer": "Session not found."}

    session_chat_type = (session.chat_type or "general").strip().lower()
    if expected_chat_type and session_chat_type != expected_chat_type:
        return None, fallback_history, None, {"answer": "Session type mismatch."}

    if expected_lesson_id is not None and session.lesson_id not in (None, expected_lesson_id):
        return None, fallback_history, None, {"answer": "Session lesson mismatch."}

    messages = (
        AiChatMessage.query
        .filter_by(session_id=session.id, user_id=user_id)
        .order_by(AiChatMessage.turn_index.asc(), AiChatMessage.id.asc())
        .all()
    )

    if edit_message_id is not None:
        try:
            edit_message_id = int(edit_message_id)
        except (TypeError, ValueError):
            return None, fallback_history, None, {"answer": "Invalid edit_message_id."}

        target = next((m for m in messages if m.id == edit_message_id and m.role == "user"), None)
        if not target:
            return None, fallback_history, None, {"answer": "Editable user message not found."}

        # Keep context only up to the edited turn (excluding the edited message itself).
        model_history = [{"role": m.role, "content": m.content} for m in messages if m.turn_index < target.turn_index]

        # Replace edited prompt and remove all following messages to regenerate from this point.
        target.content = question
        target.updated_at = datetime.utcnow()
        AiChatMessage.query.filter(
            AiChatMessage.session_id == session.id,
            AiChatMessage.user_id == user_id,
            AiChatMessage.turn_index > target.turn_index
        ).delete(synchronize_session=False)

        session.updated_at = datetime.utcnow()
        db.session.flush()

        persist_state = {
            "user_id": user_id,
            "session": session,
            "mode": "edit",
            "user_message": target,
            "assistant_turn_index": target.turn_index + 1,
        }
        return session, model_history, persist_state, None

    model_history = [{"role": m.role, "content": m.content} for m in messages]
    next_turn_index = (messages[-1].turn_index if messages else 0) + 1
    persist_state = {
        "user_id": user_id,
        "session": session,
        "mode": "new",
        "user_turn_index": next_turn_index,
    }
    return session, model_history, persist_state, None


def _persist_chat_exchange(persist_state, question, answer):
    from models import db, AiChatMessage

    if not persist_state:
        return {}

    session = persist_state["session"]
    user_id = persist_state["user_id"]

    if persist_state["mode"] == "edit":
        user_message = persist_state["user_message"]
        assistant_message = AiChatMessage(
            session_id=session.id,
            user_id=user_id,
            role="assistant",
            content=answer,
            turn_index=persist_state["assistant_turn_index"],
        )
        db.session.add(assistant_message)
    else:
        user_message = AiChatMessage(
            session_id=session.id,
            user_id=user_id,
            role="user",
            content=question,
            turn_index=persist_state["user_turn_index"],
        )
        assistant_message = AiChatMessage(
            session_id=session.id,
            user_id=user_id,
            role="assistant",
            content=answer,
            turn_index=persist_state["user_turn_index"] + 1,
        )
        db.session.add(user_message)
        db.session.add(assistant_message)

    if not (session.title or "").strip():
        session.title = _trim_title(question)
    session.updated_at = datetime.utcnow()
    db.session.commit()

    return {
        "session_id": session.id,
        "user_message_id": user_message.id,
        "assistant_message_id": assistant_message.id,
    }


# ==================== HEALTH CHECK ====================
@rag_bp.route("/health", methods=["GET"])
def health_check():
    """Health check endpoint"""
    try:
        has_groq = bool(os.environ.get("GROQ_API_KEY"))
        has_openai = bool(get_openai_api_key())
        has_docs = documents_exist()
        
        return jsonify({
            "status": "healthy",
            "groq_configured": has_groq,
            "openai_configured": has_openai,
            "documents_available": has_docs,
            "rag_initialized": vector_data is not None,
            "llm_cached": llm_cache is not None
        }), 200
    except Exception as e:
        return jsonify({"status": "unhealthy", "error": str(e)}), 500


@rag_bp.route("/admin/prompt", methods=["GET"])
@jwt_required()
def get_admin_prompt():
    if not is_admin_user():
        return jsonify({"message": "Admin access required"}), 403

    return jsonify({
        "setting_key": AI_PROMPT_KEY,
        "prompt": get_ai_system_prompt()
    }), 200


@rag_bp.route("/admin/prompt", methods=["PUT"])
@jwt_required()
def update_admin_prompt():
    if not is_admin_user():
        return jsonify({"message": "Admin access required"}), 403

    data = request.json or {}
    prompt = (data.get("prompt") or "").strip()

    if not prompt:
        return jsonify({"message": "Prompt cannot be empty"}), 400

    if len(prompt) > 12000:
        return jsonify({"message": "Prompt is too long (max 12000 chars)"}), 400

    from models import db, AppSetting

    setting = AppSetting.query.filter_by(setting_key=AI_PROMPT_KEY).first()
    if not setting:
        setting = AppSetting(setting_key=AI_PROMPT_KEY, setting_value=prompt)
        db.session.add(setting)
    else:
        setting.setting_value = prompt

    db.session.commit()

    return jsonify({
        "message": "AI prompt updated successfully",
        "prompt": setting.setting_value
    }), 200


@rag_bp.route("/admin/openai-key", methods=["GET"])
@jwt_required()
def get_admin_openai_key_status():
    if not is_admin_user():
        return jsonify({"message": "Admin access required"}), 403

    from models import AppSetting

    openai_key = get_openai_api_key()
    has_db_setting = AppSetting.query.filter_by(setting_key=OPENAI_API_KEY_SETTING).first() is not None
    return jsonify({
        "setting_key": OPENAI_API_KEY_SETTING,
        "configured": bool(openai_key and openai_key.strip()),
        "api_key": openai_key or "",
        "source": "database" if has_db_setting else "env"
    }), 200


@rag_bp.route("/admin/openai-key", methods=["PUT"])
@jwt_required()
def update_admin_openai_key():
    if not is_admin_user():
        return jsonify({"message": "Admin access required"}), 403

    data = request.json or {}
    api_key = (data.get("api_key") or "").strip()

    if not api_key:
        return jsonify({"message": "OpenAI API key cannot be empty"}), 400

    if len(api_key) < 20:
        return jsonify({"message": "OpenAI API key looks invalid"}), 400

    from models import db, AppSetting

    all_settings = AppSetting.query.filter_by(setting_key=OPENAI_API_KEY_SETTING).order_by(AppSetting.id.asc()).all()
    setting = all_settings[0] if all_settings else None
    if not setting:
        setting = AppSetting(setting_key=OPENAI_API_KEY_SETTING, setting_value=api_key)
        db.session.add(setting)
    else:
        setting.setting_value = api_key
        for duplicate in all_settings[1:]:
            db.session.delete(duplicate)

    db.session.commit()

    global llm_cache
    llm_cache = None

    return jsonify({
        "message": "OpenAI API key updated successfully",
        "configured": True
    }), 200


@rag_bp.route("/admin/openai-key/test", methods=["POST"])
@jwt_required()
def test_admin_openai_key():
    if not is_admin_user():
        return jsonify({"message": "Admin access required"}), 403

    data = request.json or {}
    api_key = (data.get("api_key") or "").strip() or (get_openai_api_key() or "").strip()

    if not api_key:
        return jsonify({"message": "OpenAI API key is not configured"}), 400

    if len(api_key) < 20:
        return jsonify({"message": "OpenAI API key looks invalid"}), 400

    try:
        from langchain_openai import ChatOpenAI
        from langchain_core.messages import HumanMessage, SystemMessage
    except ImportError:
        return jsonify({"message": "Server is missing OpenAI dependencies"}), 500

    try:
        llm = ChatOpenAI(
            openai_api_key=api_key,
            model="gpt-4.1",
            temperature=0,
            timeout=20,
            max_retries=1,
        )
        response = llm.invoke([
            SystemMessage(content="You are a connectivity test assistant. Reply in one short line."),
            HumanMessage(content="Reply with: OpenAI key is working.")
        ])
        reply = (response.content or "").strip()
        return jsonify({
            "working": True,
            "message": "OpenAI API key is working",
            "reply_preview": reply[:200]
        }), 200
    except Exception as e:
        return jsonify({
            "working": False,
            "message": "OpenAI API key test failed",
            "error": str(e)
        }), 400


def get_llm():
    global llm_cache
    if llm_cache is not None:
        return llm_cache
    
    groq_key = os.environ.get("GROQ_API_KEY")
    openai_key = get_openai_api_key()
    
    # Check for placeholder values
    if groq_key and "your_" in groq_key:
        groq_key = None
    if openai_key and "your_" in openai_key:
        openai_key = None
    
    if not groq_key and not openai_key:
        print("ERROR: Valid API Key is missing. Please add GROQ_API_KEY or OPENAI_API_KEY to .env file.")
        return "ERROR_NO_KEY"

    try:
        from langchain_openai import ChatOpenAI
        from langchain_groq import ChatGroq
    except ImportError as e:
        print(f"Import error: {e}")
        return "ERROR_IMPORT"

    _temp = _coach_llm_temperature()
    if groq_key:
        print("Using Groq LLM")
        llm_cache = ChatGroq(
            groq_api_key=groq_key,
            model="llama-3.3-70b-versatile",
            temperature=_temp,
            timeout=60,
        )
    else:
        print("Using OpenAI LLM")
        llm_cache = ChatOpenAI(
            openai_api_key=openai_key,
            # model="openai/gpt-oss-120b",
            model="gpt-4.1",
            # base_url="https://api.canopywave.io/v1",
            temperature=_temp,
            timeout=30,
            max_retries=2
        )
    
    return llm_cache


def get_qa_chain():
    global vector_data
    if vector_data is not None:
        return vector_data

    if not documents_exist():
        print("No PDF files found. Skipping RAG initialization.")
        return None

    print("\n--- Initializing RAG System ---")
    llm = get_llm()
    if llm in ("ERROR_NO_KEY", "ERROR_IMPORT"):
        return llm

    # Heavy imports inside function to avoid startup crashes
    try:
        from pypdf import PdfReader
        import numpy as np
        from sentence_transformers import SentenceTransformer
    except ImportError as e:
        print(f"Import error: {e}")
        return "ERROR_IMPORT"

    # ======== Load PDFs manually ========
    PDF_DIR = UPLOAD_FOLDER
    if not os.path.exists(PDF_DIR):
        os.makedirs(PDF_DIR)
        
    all_text = ""
    print(f"Loading PDFs from {PDF_DIR}...")
 
    files = []
    for root, dirs, filenames in os.walk(PDF_DIR):
        for name in filenames:
            if name.lower().endswith(".pdf"):
                files.append(os.path.join(root, name))
    

    if not files:
        print("No PDF files found.")
        return None

    # for file_name in files:
    #     print(f"Processing {file_name}...")
    #     try:
    #         reader = PdfReader(os.path.join(PDF_DIR, file_name))
    #         for page in reader.pages:
    #             text = page.extract_text()
    #             if text:
    #                 all_text += text + "\n"
    #     except Exception as e:
    #         print(f"Error reading {file_name}: {e}")

    for file_path in files:
        print(f"Processing {file_path}...")
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    all_text += text + "\n"
        except Exception as e:
            print(f"Error reading {file_path}: {e}")


    if not all_text.strip():
        print("No text extracted from PDFs.")
        return None

    print(f"Splitting text into chunks (Total length: {len(all_text)})...")
    chunks = [all_text[i:i+1000] for i in range(0, len(all_text), 850)]

    print(f"Generating embeddings for {len(chunks)} chunks via Local Model (SentenceTransformer)...")
    embed_model = SentenceTransformer('all-MiniLM-L6-v2')
    chunk_embeddings = embed_model.encode(chunks)

    print("RAG System Initialized.")
    vector_data = (chunks, chunk_embeddings, llm, embed_model, bool(os.environ.get("GROQ_API_KEY")))
    return vector_data

# ==================== ASK QUESTION ====================
@rag_bp.route("/ask", methods=["POST"])
@jwt_required(optional=True)
def ask():
    data = request.json or {}
    question = (data.get("question") or "").strip()
    if not question:
        return jsonify({"answer": "Please provide a question."}), 400

    persist_state = None
    try:
        _, conversation_history, persist_state, state_error = _prepare_chat_state(
            question=question,
            data=data,
            expected_chat_type="general"
        )
        if state_error:
            return jsonify(state_error), 400

        followup_anchor = _build_followup_anchor_message(conversation_history, question)

        auto_export = _try_auto_export(question, conversation_history)
        if auto_export:
            ids = _persist_chat_exchange(persist_state, question, auto_export["answer"])
            return jsonify({**auto_export, "sources_count": 0, **ids}), 200

        from langchain_core.messages import SystemMessage, HumanMessage

        if not documents_exist():
            llm = get_llm()
            if llm == "ERROR_NO_KEY":
                return jsonify({"answer": "API key is missing. Please add a valid GROQ_API_KEY or OPENAI_API_KEY and restart backend."}), 500
            if llm == "ERROR_IMPORT":
                return jsonify({"answer": "Server is missing required LLM dependencies."}), 500

            messages = [SystemMessage(content=get_ai_system_prompt())]
            if followup_anchor:
                messages.append(SystemMessage(content=followup_anchor))
            messages.extend(_build_langchain_history(conversation_history))
            messages.append(HumanMessage(content=question))

            response = llm.invoke(messages)
            answer = response.content
            ids = _persist_chat_exchange(persist_state, question, answer)
            return jsonify({"answer": answer, "sources_count": 0, **ids}), 200

        result = get_qa_chain()
        if result == "ERROR_NO_KEY":
            return jsonify({"answer": "API key is missing. Please add a valid GROQ_API_KEY or OPENAI_API_KEY and restart backend."}), 500
        if result == "ERROR_IMPORT":
            return jsonify({"answer": "Server is missing required RAG dependencies."}), 500
        if result is None:
            return jsonify({"answer": "No documents found. Please add PDFs to 'documents' folder and restart server."}), 404

        chunks, chunk_embeddings, llm, embed_model, _ = result

        context, top_idx, weak_match = _parent_library_retrieve(
            question, conversation_history, chunks, chunk_embeddings, embed_model, top_k=6
        )

        dynamic_prompt = get_ai_system_prompt()
        if weak_match:
            dynamic_prompt += (
                "\n\nNote: The uploaded library may only loosely match this question. "
                "Coach naturally from experience; you can briefly mention that a more on-topic document would sharpen answers later."
            )

        messages = [SystemMessage(content=dynamic_prompt)]
        if followup_anchor:
            messages.append(SystemMessage(content=followup_anchor))
        messages.extend(_build_langchain_history(conversation_history))
        user_block = (
            "Reference library excerpts (from the family's uploaded documents):\n"
            f"{context}\n\n---\n\nParent's message:\n{question}"
        )
        messages.append(HumanMessage(content=user_block))

        response = llm.invoke(messages)
        answer = response.content
        ids = _persist_chat_exchange(persist_state, question, answer)
        return jsonify({"answer": answer, "sources_count": len(top_idx), **ids}), 200

    except Exception as e:
        from models import db
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({"answer": f"Error: {str(e)}"}), 500


# ==================== UPLOAD PDF ====================
@rag_bp.route("/upload", methods=["POST"])
@jwt_required()
def upload_document():
    """Upload PDF document for RAG system"""
    try:
        if not is_admin_user():
            return jsonify({"message": "Admin access required"}), 403

        if 'file' not in request.files:
            return jsonify({"message": "No file provided"}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({"message": "No file selected"}), 400
        
        if not allowed_file(file.filename):
            return jsonify({"message": "Only PDF files are allowed"}), 400
        
        file_data = file.read()
        if len(file_data) > MAX_FILE_SIZE:
            return jsonify({"message": "File size exceeds 50MB limit"}), 413
        
        # Secure filename and save
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_")
        filename = timestamp + filename
        
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        with open(filepath, 'wb') as f:
            f.write(file_data)
        
        # Reset RAG system to reload documents
        global vector_data
        vector_data = None
        
        return jsonify({
            "message": "File uploaded successfully!",
            "filename": filename,
            "size": len(file_data)
        }), 201
    
    except Exception as e:
        return jsonify({"message": f"Upload error: {str(e)}"}), 500


# ==================== LIST DOCUMENTS ====================
@rag_bp.route("/documents", methods=["GET"])
@jwt_required()
def list_documents():
    """List all uploaded PDF documents"""
    try:
        if not is_admin_user():
            return jsonify({"message": "Admin access required"}), 403

        if not os.path.exists(UPLOAD_FOLDER):
            return jsonify({"documents": [], "count": 0}), 200
        
        documents = []
        for root, dirs, filenames in os.walk(UPLOAD_FOLDER):
            for filename in filenames:
                if filename.endswith(".pdf"):
                    filepath = os.path.join(root, filename)

            if filename.endswith('.pdf'):
                filepath = os.path.join(UPLOAD_FOLDER, filename)
                size = os.path.getsize(filepath)
                modified_time = os.path.getmtime(filepath)
                
                documents.append({
                    "filename": filename,
                    "size": size,
                    "size_mb": round(size / (1024 * 1024), 2),
                    "uploaded_at": datetime.fromtimestamp(modified_time).isoformat()
                })
        
        # Sort by upload time (newest first)
        documents.sort(key=lambda x: x['uploaded_at'], reverse=True)
        
        return jsonify({
            "documents": documents,
            "count": len(documents)
        }), 200
    
    except Exception as e:
        return jsonify({"message": f"Error: {str(e)}"}), 500


# ==================== DELETE DOCUMENT ====================
@rag_bp.route("/documents/<filename>", methods=["DELETE"])
@jwt_required()
def delete_document(filename):
    """Delete an uploaded PDF document"""
    try:
        if not is_admin_user():
            return jsonify({"message": "Admin access required"}), 403

        # Prevent directory traversal attacks
        if '/' in filename or '\\' in filename or filename.startswith('.'):
            return jsonify({"message": "Invalid filename"}), 400
        
        filepath = os.path.join(UPLOAD_FOLDER, secure_filename(filename))
        
        if not os.path.exists(filepath):
            return jsonify({"message": "File not found"}), 404
        
        if not filepath.endswith('.pdf'):
            return jsonify({"message": "Only PDF files can be deleted"}), 400
        
        os.remove(filepath)
        
        # Reset RAG system
        global vector_data
        vector_data = None
        
        return jsonify({"message": "Document deleted successfully!"}), 200
    
    except Exception as e:
        return jsonify({"message": f"Error: {str(e)}"}), 500


# ==================== GET RAG STATUS ====================
@rag_bp.route("/status", methods=["GET"])
@jwt_required()
def rag_status():
    """Get RAG system status"""
    try:
        if not is_admin_user():
            return jsonify({"message": "Admin access required"}), 403

        doc_count = 0
        total_size = 0
        
        if os.path.exists(UPLOAD_FOLDER):
            for root, dirs, filenames in os.walk(UPLOAD_FOLDER):
                for filename in filenames:
                    if filename.endswith(".pdf"):
                        filepath = os.path.join(root, filename)

                if filename.endswith('.pdf'):
                    doc_count += 1
                    filepath = os.path.join(UPLOAD_FOLDER, filename)
                    total_size += os.path.getsize(filepath)
        
        is_initialized = vector_data is not None
        
        return jsonify({
            "initialized": is_initialized,
            "documents_count": doc_count,
            "total_size_mb": round(total_size / (1024 * 1024), 2),
            "documents_folder": UPLOAD_FOLDER
        }), 200
    
    except Exception as e:
        return jsonify({"message": f"Error: {str(e)}"}), 500


# ==================== CREATE DOCUMENT EMBEDDINGS ====================
def create_document_embeddings(filepath):
    """Create embeddings for a newly uploaded document"""
    global vector_data
    
    try:
        print(f"Creating embeddings for: {filepath}")
        
        # Reset vector_data to force re-initialization with new document
        vector_data = None
        
        # Trigger re-initialization
        get_qa_chain()
        
        print(f"Embeddings created successfully for: {filepath}")
        return True
    
    except Exception as e:
        print(f"Error creating embeddings: {e}")
        return False


# ==================== ASK LESSON-SPECIFIC QUESTION (LLM WITH RETRIEVAL TOOL) ====================

# Cache for lesson embeddings to avoid re-processing on every question
_lesson_cache = {}

def _load_lesson_data(lesson_id):
    """Load and cache lesson text + embeddings"""
    if lesson_id in _lesson_cache:
        return _lesson_cache[lesson_id]
    
    from models import Lesson
    lesson = Lesson.query.get_or_404(lesson_id)
    lesson_file_path = os.path.join(UPLOAD_FOLDER, lesson.file_path)
    
    if not os.path.exists(lesson_file_path):
        return None
    
    text_content = ""
    
    if lesson.file_type == 'pdf':
        from pypdf import PdfReader
        reader = PdfReader(lesson_file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_content += page_text + "\n"
    elif lesson.file_type == 'txt':
        with open(lesson_file_path, 'r', encoding='utf-8') as f:
            text_content = f.read()
    
    if not text_content.strip():
        return None
    
    # Chunk the text
    chunk_size = 1000
    overlap = 150
    chunks = []
    for i in range(0, len(text_content), chunk_size - overlap):
        chunks.append(text_content[i:i + chunk_size])
    
    # Generate embeddings
    from sentence_transformers import SentenceTransformer
    embed_model = SentenceTransformer('all-MiniLM-L6-v2')
    chunk_embeddings = embed_model.encode(chunks)
    
    _lesson_cache[lesson_id] = {
        "lesson": lesson,
        "chunks": chunks,
        "embeddings": chunk_embeddings,
        "embed_model": embed_model,
        "full_text": text_content
    }
    return _lesson_cache[lesson_id]


def _retrieve_from_lesson(query, lesson_data, top_k=5):
    """Retrieval tool: search lesson chunks by semantic similarity"""
    import numpy as np
    
    chunks = lesson_data["chunks"]
    chunk_embeddings = lesson_data["embeddings"]
    embed_model = lesson_data["embed_model"]
    
    query_emb = embed_model.encode([query])[0]
    similarities = [
        float(np.dot(query_emb, ce) / (np.linalg.norm(query_emb) * np.linalg.norm(ce)))
        for ce in chunk_embeddings
    ]
    top_idx = sorted(range(len(similarities)), key=lambda i: similarities[i], reverse=True)[:top_k]
    
    results = []
    for i in top_idx:
        results.append({"chunk_index": i, "score": round(similarities[i], 4), "text": chunks[i]})
    return results


@rag_bp.route("/ask-lesson/<int:lesson_id>", methods=["POST"])
@jwt_required()
def ask_lesson(lesson_id):
    """Ask a question about a specific lesson and persist chat history."""
    data = request.json or {}
    question = (data.get("question") or "").strip()

    if not question:
        return jsonify({"answer": "Please provide a question."}), 400

    persist_state = None
    try:
        _, conversation_history, persist_state, state_error = _prepare_chat_state(
            question=question,
            data=data,
            expected_chat_type="lesson",
            expected_lesson_id=lesson_id,
        )
        if state_error and data.get("session_id") is not None:
            return jsonify(state_error), 400

        followup_anchor = _build_followup_anchor_message(conversation_history, question)

        auto_export = _try_auto_export(question, conversation_history)
        if auto_export:
            ids = _persist_chat_exchange(persist_state, question, auto_export["answer"])
            return jsonify({
                **auto_export,
                "lesson_name": None,
                "sources_count": 0,
                "search_queries": [],
                **ids,
            }), 200

        lesson_data = _load_lesson_data(lesson_id)
        if lesson_data is None:
            return jsonify({"answer": "Lesson file not found or unreadable."}), 404

        lesson = lesson_data["lesson"]

        llm = get_llm()
        if llm == "ERROR_NO_KEY":
            return jsonify({"answer": "API key is missing."}), 500
        if llm == "ERROR_IMPORT":
            return jsonify({"answer": "Server is missing required LLM dependencies."}), 500

        from langchain_core.messages import SystemMessage, HumanMessage

        recent_for_search = _build_retrieval_query_text(
            question, conversation_history, max_user_turns=2, max_chars=500
        )
        search_planning_prompt = f"""Lesson title: "{lesson.name}".

Recent chat (may include follow-ups like "what about the second part?"):
{recent_for_search}

Latest message:
{question}

Write 1-3 very short search phrases (keywords) to find the right part of the lesson PDF. One phrase per line. No numbering, no explanation."""

        search_response = llm.invoke([
            SystemMessage(
                content="You output only search phrases, one per line, to retrieve lesson text. No preamble."
            ),
            HumanMessage(content=search_planning_prompt),
        ])

        search_queries = [q.strip() for q in search_response.content.strip().split("\n") if q.strip()]
        if not search_queries:
            search_queries = [question]

        all_retrieved = []
        seen_indices = set()
        for query in search_queries[:3]:
            results = _retrieve_from_lesson(query, lesson_data, top_k=3)
            for item in results:
                if item["chunk_index"] not in seen_indices:
                    seen_indices.add(item["chunk_index"])
                    all_retrieved.append(item)

        all_retrieved.sort(key=lambda x: x["score"], reverse=True)
        top_chunks = all_retrieved[:5]

        retrieved_context = "\n\n---\n\n".join([
            f"[Section {r['chunk_index'] + 1}] (relevance: {r['score']})\n{r['text']}"
            for r in top_chunks
        ])

        system_prompt = f"""You are a warm, human-sounding tutor for the lesson "{lesson.name}".

Sound natural—like a patient teacher or parent, not a textbook. Use clear, plain language. Short questions deserve focused answers; bigger "explain this" questions can use bullets or small sections.

Use the retrieved excerpts below as your first source. If something is not there, say so honestly and say what the lesson does cover. You may add a small analogy or example when it truly helps.

Retrieved lesson excerpts:
{retrieved_context}"""

        messages = [SystemMessage(content=system_prompt)]
        if followup_anchor:
            messages.append(SystemMessage(content=followup_anchor))
        messages.extend(_build_langchain_history(conversation_history))
        messages.append(HumanMessage(content=question))

        response = llm.invoke(messages)
        answer = response.content
        ids = _persist_chat_exchange(persist_state, question, answer)

        return jsonify({
            "answer": answer,
            "lesson_name": lesson.name,
            "sources_count": len(top_chunks),
            "search_queries": search_queries,
            **ids,
        }), 200

    except Exception as e:
        from models import db
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({"answer": f"Error: {str(e)}"}), 500


@rag_bp.route("/exports/<filename>", methods=["GET"])
def download_export(filename):
    safe_name = secure_filename(filename)
    if not safe_name:
        return jsonify({"message": "Invalid filename"}), 400

    file_path = os.path.join(EXPORT_FOLDER, safe_name)
    if not os.path.exists(file_path):
        return jsonify({"message": "File not found"}), 404

    if _is_export_expired(file_path):
        try:
            os.remove(file_path)
        except OSError:
            pass
        return jsonify({"message": "This download link has expired. Please generate a new file."}), 410

    _cleanup_expired_exports()

    return send_from_directory(EXPORT_FOLDER, safe_name, as_attachment=True)
